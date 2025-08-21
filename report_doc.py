import os
import json
from datetime import datetime
from docx import Document

# Tái sử dụng hàm đọc của bạn
from data_read import read_data, read_fields   # ensure data_read.py exports these

# Đường dẫn gốc dự án (thư mục chứa file này)
PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))

# Tên file mặc định (thay đổi nếu bạn đặt tên khác)
CONFIG_FILE = os.path.join(PROJECT_DIR, "config.json")
DATA_FILE = os.path.join(PROJECT_DIR, "data_test.txt")
FIELDS_FILE = os.path.join(PROJECT_DIR, "fields.txt")


def load_and_merge_config(path):
    """
    Đọc config.json và merge tất cả entry cùng 'name' thành 1 dict mỗi hệ thống.
    Trả về dict: { system_name: merged_entry_dict }
    """
    if not os.path.exists(path):
        return {}

    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # nếu json là object -> đưa vào list để xử lý đồng nhất
    if isinstance(data, dict):
        data = [data]

    merged = {}
    for entry in data:
        name = (entry.get("name") or "").strip()
        if not name:
            continue
        if name not in merged:
            merged[name] = {}
        # merge shallow: sau cùng entry sẽ overwrite nếu cùng key
        merged[name].update(entry)

    return merged


def scan_structure(root_dir, max_lines=1000):
    """Trả về chuỗi mô tả cấu trúc thư mục (ghi ngắn gọn)."""
    lines = []
    for root, dirs, files in os.walk(root_dir):
        # ignore virtual env folder if present
        if ".venv" in root.split(os.sep) or "venv" in root.split(os.sep):
            continue
        level = root.replace(root_dir, "").count(os.sep)
        indent = '    ' * level
        folder_name = os.path.basename(root) or root
        lines.append(f"{indent}{folder_name}/")
        sub_indent = '    ' * (level + 1)
        for f in files:
            lines.append(f"{sub_indent}{f}")
        if len(lines) > max_lines:
            lines.append(f"{sub_indent}... (truncated)")
            break
    return "\n".join(lines)


def add_dict_table(doc, title, d):
    """Thêm bảng key/value cho dict d (xuống dòng nếu text dài)."""
    doc.add_heading(title, level=2)
    if not d:
        doc.add_paragraph("⚠ Không có dữ liệu")
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Key"
    hdr[1].text = "Value"
    for k, v in d.items():
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)


def add_cases_table(doc, title, cases):
    """Thêm bảng test cases (danh sách dict)."""
    doc.add_heading(title, level=2)
    if not cases:
        doc.add_paragraph("Không có test case.")
        return
    # header theo keys của first item (ổn định do read_data trả dicts cùng key)
    keys = list(cases[0].keys())
    table = doc.add_table(rows=1, cols=len(keys))
    table.style = "Table Grid"
    for i, k in enumerate(keys):
        table.rows[0].cells[i].text = k
    for case in cases:
        row = table.add_row().cells
        for i, k in enumerate(keys):
            row[i].text = str(case.get(k, ""))


def generate_report(output_dir=PROJECT_DIR):
    # --- load config, fields, data test (tái sử dụng hàm bạn có) ---
    systems = load_and_merge_config(CONFIG_FILE)  # dict: name -> merged config dict
    # read_fields your function returns dict mapping field_key->xpath (or similar)
    try:
        fields = read_fields(FIELDS_FILE)
    except Exception as e:
        fields = {}
        print(f"[WARN] Không đọc được fields từ {FIELDS_FILE}: {e}")

    # read_data: reuse your function; it will return only lines matching the data_type
    try:
        login_cases = read_data(DATA_FILE, "login")
        register_cases = read_data(DATA_FILE, "register")
    except Exception as e:
        login_cases, register_cases = [], []
        print(f"[WARN] Lỗi khi đọc data_test: {e}")

    # --- start document ---
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = f"SME_Money_Report_{ts}.docx"
    out_path = os.path.join(output_dir, out_name)
    doc = Document()

    doc.add_heading("BÁO CÁO TỰ ĐỘNG - SME Money & Sàn Tài Chính", level=0)
    doc.add_paragraph(f"Ngày tạo: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Đường dẫn dự án: {PROJECT_DIR}")
    doc.add_paragraph("Ghi chú: Sử dụng lại hàm read_fields() và read_data() hiện có.")

    # cấu trúc thư mục
    doc.add_heading("1. Cấu trúc thư mục", level=1)
    doc.add_paragraph(scan_structure(PROJECT_DIR))

    # fields chung
    doc.add_heading("2. Fields (XPath/CSS selectors)", level=1)
    if isinstance(fields, dict) and fields:
        add_dict_table(doc, "Fields mapping", fields)
    else:
        doc.add_paragraph(f"Không đọc được fields từ {FIELDS_FILE} hoặc file rỗng.")

    # testcases chung (tách login/register)
    doc.add_heading("3. Test cases (tách Login & Register)", level=1)
    add_cases_table(doc, "3.1 Login test cases", login_cases)
    add_cases_table(doc, "3.2 Register test cases", register_cases)

    # per-system sections
    doc.add_heading("4. Thông tin theo hệ thống", level=1)
    if not systems:
        doc.add_paragraph(f"Không tìm thấy config hoặc {CONFIG_FILE} rỗng.")
    else:
        for sys_name, cfg in systems.items():
            doc.add_heading(sys_name, level=2)
            # show merged config keys (url, url_register, db, ...)
            add_dict_table(doc, "Cấu hình hệ thống", cfg)

            # show which testcases are relevant — we don't filter testcases by system
            # (user requested same file; so we will show both features for each system)
            doc.add_paragraph("Tính năng: Đăng nhập (Login)")
            add_cases_table(doc, "Login testcases (áp dụng chung)", login_cases)

            doc.add_paragraph("Tính năng: Đăng ký (Register)")
            add_cases_table(doc, "Register testcases (áp dụng chung)", register_cases)
    doc.add_heading("5. Luồng chạy script", level=1)
    doc.add_paragraph(
        "1. Đọc config từ file config.json\n"
        "2. Đọc dữ liệu test từ file data_test.txt\n"
        "3. Khởi tạo driver Selenium\n"
        "4. Thực thi các test case\n"
        "5. Ghi log vào result.txt hoặc result_register.txt\n"
        "6. Đóng driver và kết thúc"
    )
    # Lưu file
    doc.save(out_path)
    print(f"[OK] Báo cáo đã tạo: {out_path}")
    return out_path


if __name__ == "__main__":
    generate_report()
